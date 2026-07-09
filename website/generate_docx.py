from docx import Document
from io import BytesIO
from flask import send_file
import re


def create_docx_from_markdown(markdown_content, document_type, job_title):
    
    doc = Document()

    doc.add_heading(document_type, level=0)
    doc.add_heading(job_title, level=1)

    for line in markdown_content.splitlines():
        line = line.strip()

        if not line:
            continue

        if line.startswith("**") or line.startswith("##"):
            doc.add_heading(line[2:], level=1)
        
        elif line.startswith("***") or line.startswith("###"):
            doc.add_heading(line[3:], level=2)
        
        elif line.startswith("*") or line.startswith("-"):
            doc.add_paragraph(line[1:], style='List Bullet')
        
        else:
            doc.add_paragraph(line)
        
    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_job_title = re.sub(r'[<>:"/\\|?*]', "_", job_title)

    filename = f"{document_type}_{safe_job_title.replace(' ', '_')}.docx"


    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',

    )